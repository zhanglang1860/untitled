#!/usr/bin/env python
# -*- coding: utf-8 -*-


# import pandas as pd
import docx  # 处理word文档
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 引入文字的位置样式
from docx.shared import Pt  # 设置字体大小
from docx.oxml.ns import qn


def read_docx(path):
    data = docx.Document(path)
    #   for index, para in enumerate(data.paragraphs):
    #       print(index, para.text)
    #   p = data.paragraphs[1]
    #   print(p.text)
    #   for p in data.paragraphs:
    #       print(p.text)

    # p = data.add_paragraph(u'第一段', style=None)
    # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # data.add_paragraph(u'第二段', style=None)
    # data.save(r'C:\Users\Administrator\PycharmProjects\untitled\files\test02.docx')

    paragraph1 = data.add_paragraph('关于的批复')  # 用add_paragraph添加新段落
    #paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 位置居中
    paragraph1.style = 'asdas'  # 段落样式
    # run1 = paragraph1.add_run('关于同意某某医学院\n附属某某医院第九次代表大会代表候选人的批复')  # 用add_run添加更多的使用

    # run1.font.size = Pt(18)  # 字体大小
    # run1.font.name = u'黑体'  # 字体样式
    # run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')  # 字体样式
    # run1.font.bold = True  # 字体加粗
    # paragraph1.space_after = Pt(10)  # 段前间距
    # paragraph1.paragraph_format.line_spacing = Pt(32)  # 段内间距

    data.save(r'C:\Users\Administrator\PycharmProjects\untitled\files\test03.docx')
if __name__ == "__main__":
    read_docx(r'C:\Users\Administrator\PycharmProjects\untitled\files\test03.docx')
